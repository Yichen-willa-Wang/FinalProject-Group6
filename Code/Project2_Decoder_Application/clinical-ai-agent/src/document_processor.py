"""
Enhanced Document Processor
Supports PDF, Word (docx), with smart chunking for tables and images
"""

import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from pathlib import Path
import cv2
import numpy as np
from typing import List, Dict, Tuple
import pdfplumber
import pandas as pd

# For Word documents
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed, Word support disabled")

class DocumentProcessor:
    """Process PDF and Word documents with smart content extraction"""
    
    def __init__(self):
        self.tesseract_config = '--psm 6 --oem 3'
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def process_document(self, file_path: str) -> Dict:
        """
        Main entry point - auto-detect format and process
        
        Returns:
            {
                'text_content': [...],
                'tables': [...],
                'images': [...]
            }
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.process_pdf(str(file_path))
        elif suffix in ['.docx', '.doc']:
            return self.process_word(str(file_path))
        else:
            raise ValueError(f"Unsupported format: {suffix}")
    
    def process_pdf(self, pdf_path: str) -> Dict:
        """Process PDF with text, tables, and images"""
        print(f"Processing PDF: {pdf_path}")
        
        result = {
            'text_content': [],
            'tables': [],
            'images': []
        }
        
        # Extract text
        result['text_content'] = self._extract_pdf_text(pdf_path)
        
        # Extract tables
        result['tables'] = self._extract_pdf_tables(pdf_path)
        
        # Extract images
        result['images'] = self._extract_pdf_images(pdf_path)
        
        return result
    
    def _extract_pdf_text(self, pdf_path: str) -> List[Dict]:
        """Extract text from PDF"""
        doc = fitz.open(pdf_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            # Check if scanned (low text content)
            if len(text.strip()) < 100:
                # Use OCR
                pix = page.get_pixmap(dpi=300)
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                    pix.height, pix.width, pix.n
                )
                if pix.n == 4:  # RGBA
                    img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
                
                gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
                _, binary = cv2.threshold(
                    gray, 0, 255, 
                    cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )
                
                text = pytesseract.image_to_string(
                    binary,
                    lang='chi_sim+chi_tra',
                    config=self.tesseract_config
                )
            
            text_content.append({
                'page': page_num + 1,
                'text': text,
                'type': 'text',
                'char_count': len(text)
            })
        
        doc.close()
        return text_content
    
    def _extract_pdf_tables(self, pdf_path: str) -> List[Dict]:
        """Extract tables from PDF using pdfplumber"""
        tables_data = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    
                    for table_idx, table in enumerate(tables):
                        if table and len(table) > 0:
                            # Convert to DataFrame
                            df = pd.DataFrame(table[1:], columns=table[0])
                            
                            # Convert table to text description
                            table_text = self._table_to_text(df)
                            
                            tables_data.append({
                                'page': page_num + 1,
                                'table_index': table_idx + 1,
                                'dataframe': df,
                                'text_representation': table_text,
                                'rows': len(df),
                                'columns': len(df.columns)
                            })
        except Exception as e:
            print(f"Table extraction warning: {e}")
        
        return tables_data
    
    def _extract_pdf_images(self, pdf_path: str) -> List[Dict]:
        """Extract images from PDF"""
        doc = fitz.open(pdf_path)
        images_info = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    
                    images_info.append({
                        'page': page_num + 1,
                        'image_index': img_index + 1,
                        'width': base_image.get('width', 0),
                        'height': base_image.get('height', 0),
                        'ext': base_image.get('ext', 'png')
                    })
                except:
                    continue
        
        doc.close()
        return images_info
    
    def process_word(self, docx_path: str) -> Dict:
        """Process Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")
        
        print(f"Processing Word: {docx_path}")
        
        doc = DocxDocument(docx_path)
        
        result = {
            'text_content': [],
            'tables': [],
            'images': []
        }
        
        # Extract paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        result['text_content'] = [{
            'page': 1,
            'text': '\n'.join(full_text),
            'type': 'text',
            'char_count': len('\n'.join(full_text))
        }]
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            
            if data:
                df = pd.DataFrame(data[1:], columns=data[0])
                table_text = self._table_to_text(df)
                
                result['tables'].append({
                    'page': 1,
                    'table_index': table_idx + 1,
                    'dataframe': df,
                    'text_representation': table_text,
                    'rows': len(df),
                    'columns': len(df.columns)
                })
        
        return result
    
    def _table_to_text(self, df: pd.DataFrame) -> str:
        """Convert table to natural language text"""
        text_parts = []
        
        # Add column headers
        text_parts.append(f"Table with {len(df)} rows and {len(df.columns)} columns.")
        text_parts.append(f"Columns: {', '.join(df.columns)}")
        
        # Add data rows (limit to first 5 for chunking)
        for idx, row in df.head(5).iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            text_parts.append(row_text)
        
        if len(df) > 5:
            text_parts.append(f"... and {len(df) - 5} more rows")
        
        return "\n".join(text_parts)
